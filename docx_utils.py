from pathlib import Path
from datetime import datetime
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

BLUE = RGBColor(30, 58, 138)
BLUE2 = RGBColor(37, 99, 235)
MUTED = RGBColor(100, 116, 139)


def clean_filename(text):
    text = re.sub(r'[^a-zA-Z0-9_-]+', '_', text.strip())
    return text[:60] or 'plano_de_aula'


def parse_sections(markdown_text):
    sections = []
    title = 'Conteúdo'
    lines = []
    for raw in markdown_text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith('# '):
            continue
        if line.startswith('## '):
            if lines:
                sections.append((title, lines))
            title = line.replace('## ', '').strip()
            lines = []
        else:
            lines.append(line.replace('**', '').replace('* ', '- '))
    if lines:
        sections.append((title, lines))
    return sections


def add_header_run(paragraph, text, size=16, color=BLUE):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def generate_docx(title, content, meta=None, logo_path='assets/logo.png'):
    meta = meta or {}
    out = Path('exports')
    out.mkdir(exist_ok=True)
    path = out / f"{clean_filename(title)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal'].font.size = Pt(10.5)

    if Path(logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(logo_path, width=Inches(2.2))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_header_run(p, 'Plano de Aula Gerado com Apoio Pedagógico Inteligente', 16, BLUE)

    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    rows = [
        ('Professor', meta.get('teacher', 'Não informado'), 'Escola', meta.get('school', 'Não informada')),
        ('Componente', meta.get('subject', 'Não informado'), 'Ano/Série', meta.get('grade', 'Não informado')),
        ('Tema', meta.get('theme', 'Não informado'), 'Geração', datetime.now().strftime('%d/%m/%Y %H:%M')),
    ]
    for row_cells, data in zip(table.rows, rows):
        for cell, value in zip(row_cells.cells, data):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph('')

    for section, lines in parse_sections(content):
        p = doc.add_paragraph()
        add_header_run(p, section.upper(), 13, BLUE2)
        for line in lines:
            p = doc.add_paragraph()
            line = line.replace('- ', '').strip()
            if ':' in line and len(line.split(':',1)[0]) < 35:
                k, v = line.split(':',1)
                r = p.add_run(k.strip() + ': ')
                r.bold = True
                r.font.color.rgb = BLUE
                p.add_run(v.strip())
            else:
                p.add_run(line)

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Desenvolvido por Lucas Lopes • ADS EAD • UNIFECAF')
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED

    doc.save(path)
    return str(path)
